# файл: randomize_word_text.py

import random
import os
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import tkinter as tk
from tkinter import filedialog, messagebox


# --- Константы "естественности" ---

# Межстрочный интервал (multiple)
BASE_LINE_SPACING = 1.15        # можно поменять на 1.0, если нужно
LINE_SPACING_DELTA = 0.07       # разброс ±0.07

# Двойные / тройные пробелы между словами
EXTRA_SPACE_PROB_DOUBLE = 0.02   # 2% пробелов станут двойными
EXTRA_SPACE_PROB_TRIPLE = 0.005  # 0.5% станут тройными

# Шанс странных верхних/нижних индексов
BASELINE_WEIRD_PROB = 0.003      # 0.3% символов (sup/subscript)

# Шанс странной ширины символов (w-scale)
CHAR_SCALE_PROB = 0.003          # 0.3% символов
CHAR_SCALE_MIN = 97              # 97% ширины
CHAR_SCALE_MAX = 103             # 103% ширины

# Параметры "волны" baseline (в pt)
WAVE_STEP_MIN = 0.05   # минимальный шаг изменения baseline (pt) за символ
WAVE_STEP_MAX = 0.15   # максимальный шаг
WAVE_AMPLITUDE = 1.2   # максимальный |offset| в pt (примерно как писал)


def detect_base_font_size(doc, fallback=14):
    """
    Пытаемся определить базовый размер шрифта по первому абзацу/рану.
    Если не нашли - используем fallback (по умолчанию 14 pt).
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is not None:
                try:
                    return run.font.size.pt
                except Exception:
                    pass
    # Если в раннах не найдено - можно попробовать стиль абзаца
    for paragraph in doc.paragraphs:
        if paragraph.style and paragraph.style.font and paragraph.style.font.size:
            try:
                return paragraph.style.font.size.pt
            except Exception:
                pass
    return fallback


def clear_paragraph(paragraph):
    """
    Полностью очищает содержимое параграфа (удаляет все run'ы).
    """
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def get_or_add_rPr(run):
    """
    Возвращает или создаёт <w:rPr> для данного run через низкоуровневый XML.
    """
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    return rPr


def add_random_extra_spaces(text,
                            prob_double=EXTRA_SPACE_PROB_DOUBLE,
                            prob_triple=EXTRA_SPACE_PROB_TRIPLE):
    """
    В небольшом проценте случаев заменяет " " на "  " или "   ".
    """
    result_parts = []
    for ch in text:
        if ch == " ":
            r = random.random()
            if r < prob_triple:
                result_parts.append("   ")
            elif r < prob_triple + prob_double:
                result_parts.append("  ")
            else:
                result_parts.append(" ")
        else:
            result_parts.append(ch)
    return "".join(result_parts)


def apply_wave_baseline(run, baseline_offset_pt):
    """
    Применяет плавное смещение baseline к run через w:position.
    baseline_offset_pt - в пунктах (pt).
    В Word w:position хранится в half-points, поэтому умножаем на 2.
    """
    rPr = get_or_add_rPr(run)
    pos_el = OxmlElement('w:position')
    # конвертируем pt в half-points
    val = int(baseline_offset_pt * 2)
    pos_el.set(qn('w:val'), str(val))
    rPr.append(pos_el)


def apply_random_run_effects(run,
                             baseline_prob=BASELINE_WEIRD_PROB,
                             scale_prob=CHAR_SCALE_PROB,
                             scale_min=CHAR_SCALE_MIN,
                             scale_max=CHAR_SCALE_MAX):
    """
    Редкие «грязные» эффекты:
      - иногда символ становится верхним/нижним индексом (w:vertAlign)
      - иногда меняется ширина символов (w:w 97–103)

    ВАЖНО: здесь больше НЕ трогаем w:position,
    чтобы не ломать красивую плавную волну.
    """
    text = run.text
    if not text or text.isspace():
        return

    rPr = get_or_add_rPr(run)

    # 1) Случайный верхний/нижний индекс
    if random.random() < baseline_prob:
        vert_el = OxmlElement('w:vertAlign')
        vert_el.set(qn('w:val'), random.choice(['superscript', 'subscript']))
        rPr.append(vert_el)

    # 2) Рандомная ширина символов (CharacterScale, w:w)
    if random.random() < scale_prob:
        scale_el = OxmlElement('w:w')
        scale = random.randint(scale_min, scale_max)
        scale_el.set(qn('w:val'), str(scale))
        rPr.append(scale_el)


def randomize_paragraph_text(
    paragraph,
    base_size_pt=14,
    delta_pt=1,
    base_line_spacing=BASE_LINE_SPACING,
    line_spacing_delta=LINE_SPACING_DELTA
):
    """
    1) Добавляет в начало абзаца от 1 до 3 пробелов.
    2) Пересоздаёт текст абзаца посимвольно с размером шрифта base_size_pt ± delta_pt.
    3) Чуть рандомизирует межстрочный интервал (line_spacing).
    4) Добавляет редкие двойные/тройные пробелы.
    5) Для каждого абзаца создаёт плавную "волну" baseline по всей строке.
    """
    # Рандомный межстрочный интервал (multiple)
    try:
        pf = paragraph.paragraph_format
        ls = base_line_spacing + random.uniform(-line_spacing_delta, line_spacing_delta)
        pf.line_spacing = ls
    except Exception:
        pass

    original_text = paragraph.text

    # Пустой / только пробелы
    if not original_text.strip():
        clear_paragraph(paragraph)
        spaces = " " * random.randint(1, 3)
        run = paragraph.add_run(spaces)
        run.font.size = Pt(base_size_pt)
        return

    # Непустой абзац
    clear_paragraph(paragraph)

    # 1–3 пробела в начале
    leading_spaces = " " * random.randint(1, 3)
    full_text = leading_spaces + original_text.lstrip()

    # Редкие двойные/тройные пробелы
    full_text = add_random_extra_spaces(full_text)

    # --- Параметры волны baseline для этого абзаца ---
    baseline_offset = 0.0
    direction = random.choice([-1, 1])  # стартуем либо вверх, либо вниз

    # Посимвольная генерация
    for ch in full_text:
        size = random.randint(
            int(base_size_pt - delta_pt),
            int(base_size_pt + delta_pt),
        )
        run = paragraph.add_run(ch)
        run.font.size = Pt(size)

        if not ch.isspace():
            # Плавно двигаем baseline
            step = random.uniform(WAVE_STEP_MIN, WAVE_STEP_MAX)
            baseline_offset += direction * step

            # Меняем направление, если вышли за амплитуду
            if baseline_offset > WAVE_AMPLITUDE:
                direction = -1
            elif baseline_offset < -WAVE_AMPLITUDE:
                direction = 1

            # Применяем смещение baseline к текущему run
            apply_wave_baseline(run, baseline_offset)

            # Дополнительные редкие эффекты (sup/sub, ширина)
            apply_random_run_effects(run)


def randomize_doc(doc, base_size_pt=14, delta_pt=1):
    """
    Применяем рандомизацию ко всем абзацам и тексту в таблицах.
    """
    # Обычные абзацы
    for paragraph in doc.paragraphs:
        randomize_paragraph_text(paragraph, base_size_pt, delta_pt)

    # Текст в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    randomize_paragraph_text(paragraph, base_size_pt, delta_pt)


def main():
    root = tk.Tk()
    root.withdraw()

    messagebox.showinfo(
        "Рандомизация Word",
        "Выберите .docx файл. Приложение добавит:\n"
        "- случайные отступы в начале строк,\n"
        "- разброс размера шрифта (±1 pt),\n"
        "- лёгкую неровность межстрочных интервалов,\n"
        "- редкие двойные/тройные пробелы,\n"
        "- плавную «волну» baseline по строке,\n"
        "- микродефекты символов (верхние/нижние индексы, ширина)."
    )

    file_path = filedialog.askopenfilename(
        title="Выберите Word-файл",
        filetypes=[("Word документ", "*.docx")]
    )

    if not file_path:
        messagebox.showwarning("Отмена", "Файл не выбран. Работа завершена.")
        return

    try:
        doc = Document(file_path)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось открыть файл:\n{e}")
        return

    # Определяем базовый размер шрифта по документу (или 14 pt по умолчанию)
    base_size_pt = detect_base_font_size(doc, fallback=14)

    # Основная рандомизация
    randomize_doc(doc, base_size_pt=base_size_pt, delta_pt=1)

    # Новое имя файла
    folder, name = os.path.split(file_path)
    base, ext = os.path.splitext(name)
    new_name = f"{base}_random_font{ext}"
    new_path = os.path.join(folder, new_name)

    try:
        doc.save(new_path)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить файл:\n{e}")
        return

    messagebox.showinfo(
        "Готово",
        f"Новый файл сохранён как:\n{new_path}\n\n"
        f"Базовый размер шрифта ~ {base_size_pt} pt, разброс ±1 pt.\n"
        f"Эффекты: волна baseline, неровные интервалы, пробелы и микродефекты символов."
    )


if __name__ == "__main__":
    main()
